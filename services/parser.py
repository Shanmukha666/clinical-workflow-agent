"""
Production Document Parser for Clinical Data
Supports PDF, DOCX, HTML, JSON, and structured text extraction
"""

from __future__ import annotations

import os
import re
import json
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from datetime import datetime
import asyncio
from concurrent.futures import ThreadPoolExecutor
import hashlib

# PDF parsing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# HTML parsing
try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ==================== Configuration ====================

class ParserConfig:
    """Parser configuration settings"""
    
    # Supported formats
    SUPPORTED_FORMATS = {
        '.txt': 'text',
        '.json': 'json',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.html': 'html',
        '.htm': 'html',
        '.xml': 'xml'
    }
    
    # PDF parsing settings
    PDF_USE_PDFPLUMBER = os.getenv("PDF_USE_PDFPLUMBER", "True").lower() == "true"
    PDF_EXTRACT_TABLES = os.getenv("PDF_EXTRACT_TABLES", "True").lower() == "true"
    
    # Text cleaning
    REMOVE_EXTRA_WHITESPACE = True
    NORMALIZE_LINE_BREAKS = True
    REMOVE_SPECIAL_CHARS = False  # Keep clinical notation
    
    # Cache settings
    CACHE_PARSED_RESULTS = os.getenv("CACHE_PARSER", "True").lower() == "true"
    CACHE_DIR = Path("/tmp/parser_cache")
    
    # Max file size (50MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024


class ParseResult:
    """Parsed document result"""
    
    def __init__(self, text: str, metadata: Dict[str, Any], format: str, 
                 confidence: float = 1.0, tables: Optional[List[Dict]] = None):
        self.text = text
        self.metadata = metadata
        self.format = format
        self.confidence = confidence
        self.tables = tables or []
        self.processing_time = 0.0
        self.word_count = len(text.split())
        self.char_count = len(text)
        self.timestamp = datetime.now().isoformat()
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "metadata": self.metadata,
            "format": self.format,
            "confidence": self.confidence,
            "tables": self.tables,
            "processing_time_ms": self.processing_time,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "timestamp": self.timestamp
        }


# ==================== Format-Specific Parsers ====================

class PDFParser:
    """PDF document parser"""
    
    @staticmethod
    def parse(file_path: str) -> Tuple[str, Dict[str, Any], List[Dict]]:
        """Extract text and tables from PDF"""
        text = ""
        tables = []
        metadata = {
            "pages": 0,
            "has_tables": False,
            "encrypted": False
        }
        
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing libraries not installed")
        
        # Try pdfplumber first (better table extraction)
        if ParserConfig.PDF_USE_PDFPLUMBER:
            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata["pages"] = len(pdf.pages)
                    metadata["encrypted"] = pdf.metadata.get("Encrypted", False)
                    
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                        
                        if ParserConfig.PDF_EXTRACT_TABLES:
                            page_tables = page.extract_tables()
                            for table in page_tables:
                                if table and len(table) > 1:
                                    tables.append({
                                        "page": page_num + 1,
                                        "rows": len(table),
                                        "columns": len(table[0]) if table else 0,
                                        "data": table[:10]  # Limit stored data
                                    })
                                    metadata["has_tables"] = True
                
                return text, metadata, tables
                
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}, falling back to PyPDF2")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata["pages"] = len(reader.pages)
                metadata["encrypted"] = reader.is_encrypted
                
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
            return text, metadata, tables
            
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise


class DOCXParser:
    """Word document parser"""
    
    @staticmethod
    def parse(file_path: str) -> Tuple[str, Dict[str, Any], List[Dict]]:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing libraries not installed")
        
        text = ""
        tables = []
        metadata = {
            "paragraphs": 0,
            "has_tables": False
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    metadata["paragraphs"] += 1
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append({
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0,
                        "data": table_data[:10]
                    })
                    metadata["has_tables"] = True
            
            # Extract core properties
            core_props = doc.core_properties
            metadata.update({
                "author": core_props.author,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "title": core_props.title
            })
            
            return text, metadata, tables
            
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise


class HTMLParser:
    """HTML document parser"""
    
    @staticmethod
    def parse(file_path: str) -> Tuple[str, Dict[str, Any], List[Dict]]:
        """Extract text from HTML"""
        if not HTML_AVAILABLE:
            raise ImportError("HTML parsing libraries not installed")
        
        text = ""
        tables = []
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Get title
            title = soup.find('title')
            if title:
                metadata['title'] = title.text
            
            # Extract main content (remove scripts, styles)
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Extract tables
            for table in soup.find_all('table'):
                rows = []
                for row in table.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    if cells:
                        rows.append(cells)
                if rows:
                    tables.append({
                        "rows": len(rows),
                        "columns": len(rows[0]) if rows else 0,
                        "data": rows[:10]
                    })
                    metadata['has_tables'] = True
            
            return text, metadata, tables
            
        except Exception as e:
            logger.error(f"HTML parsing failed: {e}")
            raise


class JSONParser:
    """JSON document parser"""
    
    @staticmethod
    def parse(file_path: str) -> Tuple[str, Dict[str, Any], List[Dict]]:
        """Extract text from JSON"""
        text = ""
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            metadata["keys"] = list(data.keys()) if isinstance(data, dict) else None
            metadata["type"] = type(data).__name__
            
            # Extract text from JSON structure
            text = JSONParser._extract_text_from_json(data)
            
            return text, metadata, []
            
        except Exception as e:
            logger.error(f"JSON parsing failed: {e}")
            raise
    
    @staticmethod
    def _extract_text_from_json(obj, depth=0) -> str:
        """Recursively extract text from JSON"""
        text = []
        
        if isinstance(obj, dict):
            for key, value in obj.items():
                if isinstance(value, (str, int, float, bool)):
                    text.append(f"{key}: {value}")
                elif isinstance(value, (dict, list)):
                    text.append(JSONParser._extract_text_from_json(value, depth + 1))
        elif isinstance(obj, list):
            for item in obj:
                text.append(JSONParser._extract_text_from_json(item, depth + 1))
        elif isinstance(obj, (str, int, float, bool)):
            text.append(str(obj))
        
        return "\n".join(text)


class TextParser:
    """Plain text parser"""
    
    @staticmethod
    def parse(file_path: str) -> Tuple[str, Dict[str, Any], List[Dict]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                "size_bytes": os.path.getsize(file_path),
                "lines": len(text.splitlines()),
                "encoding": "utf-8"
            }
            
            return text, metadata, []
            
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            
            metadata = {
                "size_bytes": os.path.getsize(file_path),
                "lines": len(text.splitlines()),
                "encoding": "latin-1"
            }
            
            return text, metadata, []


# ==================== Main Parser Service ====================

class DocumentParser:
    """
    Main document parser service with format detection and caching
    """
    
    def __init__(self):
        self.cache_dir = ParserConfig.CACHE_DIR
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.executor = ThreadPoolExecutor(max_workers=4)
        
        # Register parsers
        self.parsers = {
            'pdf': PDFParser,
            'docx': DOCXParser,
            'html': HTMLParser,
            'json': JSONParser,
            'text': TextParser
        }
        
        logger.info(f"Document Parser initialized. Supported formats: {list(ParserConfig.SUPPORTED_FORMATS.keys())}")
    
    async def parse(self, file_path: str) -> ParseResult:
        """
        Parse document with automatic format detection
        """
        import time
        start_time = time.time()
        
        # Validate file
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_size = os.path.getsize(file_path)
        if file_size > ParserConfig.MAX_FILE_SIZE:
            raise ValueError(f"File too large: {file_size} bytes. Max: {ParserConfig.MAX_FILE_SIZE}")
        
        # Detect format
        file_ext = Path(file_path).suffix.lower()
        format_type = ParserConfig.SUPPORTED_FORMATS.get(file_ext, 'text')
        
        # Check cache
        cache_key = self._get_cache_key(file_path)
        if ParserConfig.CACHE_PARSED_RESULTS:
            cached = self._get_cached(cache_key)
            if cached:
                logger.info(f"Parser cache hit for {file_path}")
                cached.processing_time = (time.time() - start_time) * 1000
                return cached
        
        # Get parser
        parser_class = self.parsers.get(format_type)
        if not parser_class:
            raise ValueError(f"No parser for format: {format_type}")
        
        # Parse document
        try:
            text, metadata, tables = await self._run_parser(parser_class, file_path)
            
            # Clean text
            text = self._clean_text(text)
            
            # Create result
            result = ParseResult(
                text=text,
                metadata=metadata,
                format=format_type,
                tables=tables
            )
            result.processing_time = (time.time() - start_time) * 1000
            
            # Cache result
            if ParserConfig.CACHE_PARSED_RESULTS:
                self._cache_result(cache_key, result)
            
            logger.info(f"Parsed {file_path}: format={format_type}, words={result.word_count}, time={result.processing_time:.2f}ms")
            
            return result
            
        except Exception as e:
            logger.error(f"Parsing failed for {file_path}: {e}")
            raise
    
    async def _run_parser(self, parser_class, file_path: str) -> Tuple[str, Dict, List]:
        """Run parser in thread pool"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self.executor, parser_class.parse, file_path
        )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if ParserConfig.REMOVE_EXTRA_WHITESPACE:
            text = re.sub(r'\s+', ' ', text)
        
        if ParserConfig.NORMALIZE_LINE_BREAKS:
            text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        if ParserConfig.REMOVE_SPECIAL_CHARS:
            text = re.sub(r'[^\w\s\-:.,/()%]', '', text)
        
        return text.strip()
    
    def _get_cache_key(self, file_path: str) -> str:
        """Generate cache key from file content"""
        with open(file_path, 'rb') as f:
            content = f.read()
        return hashlib.md5(content).hexdigest()
    
    def _get_cached(self, key: str) -> Optional[ParseResult]:
        """Get cached parse result"""
        cache_file = self.cache_dir / f"{key}.json"
        if cache_file.exists():
            try:
                with open(cache_file, 'r') as f:
                    data = json.load(f)
                result = ParseResult(
                    text=data['text'],
                    metadata=data['metadata'],
                    format=data['format'],
                    confidence=data['confidence'],
                    tables=data.get('tables', [])
                )
                return result
            except Exception as e:
                logger.warning(f"Cache read failed: {e}")
        return None
    
    def _cache_result(self, key: str, result: ParseResult):
        """Cache parse result"""
        cache_file = self.cache_dir / f"{key}.json"
        try:
            with open(cache_file, 'w') as f:
                json.dump(result.to_dict(), f, default=str)
        except Exception as e:
            logger.warning(f"Cache write failed: {e}")


# ==================== Singleton Instance ====================

_parser_service: Optional[DocumentParser] = None


def get_parser_service() -> DocumentParser:
    """Get singleton parser service instance"""
    global _parser_service
    if _parser_service is None:
        _parser_service = DocumentParser()
    return _parser_service


async def extract_text_from_pdf(file_path: str) -> str:
    """Async wrapper for PDF extraction"""
    service = get_parser_service()
    result = await service.parse(file_path)
    return result.text


# ==================== Test Block ====================

if __name__ == "__main__":
    import asyncio
    
    async def test():
        print("🚀 Testing Document Parser...\n")
        
        service = get_parser_service()
        
        test_file = "sample.pdf"  # Change this
        if os.path.exists(test_file):
            result = await service.parse(test_file)
            
            print(f"✅ Parse Result:")
            print(f"   Format: {result.format}")
            print(f"   Confidence: {result.confidence:.2%}")
            print(f"   Words: {result.word_count}")
            print(f"   Time: {result.processing_time:.2f}ms")
            print(f"   Tables: {len(result.tables)}")
            print(f"\n📝 Extracted Text:\n{result.text[:500]}")
        else:
            print(f"❌ Test file not found: {test_file}")
    
    asyncio.run(test())